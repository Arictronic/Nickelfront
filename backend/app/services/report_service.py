"""
Сервис генерации и экспорта отчётов.

Поддерживаемые форматы:
- PDF (через ReportLab)
- DOCX (через python-docx)
"""

from datetime import datetime
from pathlib import Path
from typing import Optional, BinaryIO
import io
import sys

from pathlib import Path

# Добавляем корень проекта в PATH
ROOT_DIR = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(ROOT_DIR))
sys.path.insert(0, str(ROOT_DIR / "backend"))
sys.path.insert(0, str(ROOT_DIR / "analytics"))

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class PaperReportData:
    """Данные для отчёта по статье."""
    
    def __init__(self, paper: dict, metrics: dict = None):
        self.paper = paper
        self.metrics = metrics or {}
        
        # Основная информация
        self.id = paper.get("id", "N/A")
        self.title = paper.get("title", "Без названия")
        self.authors = paper.get("authors", [])
        self.journal = paper.get("journal", "N/A")
        self.publication_date = paper.get("publication_date", "N/A")
        self.doi = paper.get("doi", "N/A")
        self.source = paper.get("source", "N/A")
        self.url = paper.get("url", "")
        
        # Контент
        self.abstract = paper.get("abstract", "")
        self.full_text = paper.get("full_text", "")
        self.keywords = paper.get("keywords", [])
        
        # Метрики
        self.abstract_length = len(self.abstract) if self.abstract else 0
        self.full_text_length = len(self.full_text) if self.full_text else 0
        self.keywords_count = len(self.keywords) if self.keywords else 0
        
        # Оценка качества
        self.quality_score = self._calculate_quality_score()
        
    def _calculate_quality_score(self) -> float:
        """Вычислить оценку качества статьи."""
        score = 0
        max_score = 100
        
        if self.abstract and len(self.abstract) > 100:
            score += 20
        if self.full_text and len(self.full_text) > 1000:
            score += 30
        if self.keywords and len(self.keywords) >= 5:
            score += 20
        if self.doi and self.doi != "N/A":
            score += 15
        if self.authors and len(self.authors) > 0:
            score += 15
            
        return min(score, max_score)
    
    def get_recommendations(self) -> list[str]:
        """Получить рекомендации по улучшению."""
        recommendations = []
        
        if not self.abstract or len(self.abstract) <= 100:
            recommendations.append("Добавить или расширить аннотацию")
        if not self.full_text or len(self.full_text) <= 1000:
            recommendations.append("Добавить полный текст статьи")
        if not self.keywords or len(self.keywords) < 5:
            recommendations.append(f"Добавить ключевые слова (текущее: {self.keywords_count})")
        if not self.doi or self.doi == "N/A":
            recommendations.append("Добавить DOI")
        if not self.authors:
            recommendations.append("Добавить авторов")
            
        return recommendations


class ReportExporter:
    """Экспортёр отчётов в различные форматы."""
    
    def __init__(self, report_data: PaperReportData):
        self.report_data = report_data
        
    def export_pdf(self) -> bytes:
        """
        Экспортировать отчёт в PDF.
        
        Returns:
            Bytes PDF файла
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab не установлен. Установите: pip install reportlab")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        elements = []
        styles = getSampleStyleSheet()
        
        # Заголовок
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1e293b'),
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        elements.append(Paragraph("ОТЧЁТ ПО СТАТЬЕ", title_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Основная информация
        info_style = ParagraphStyle(
            'InfoStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#64748b'),
            spaceAfter=6
        )
        
        elements.append(Paragraph(f"<b>ID:</b> {self.report_data.id}", info_style))
        elements.append(Paragraph(f"<b>Источник:</b> {self.report_data.source}", info_style))
        elements.append(Paragraph(f"<b>Дата генерации:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}", info_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Название
        title_para = Paragraph(f"<b>{self.report_data.title}</b>", styles['Heading2'])
        elements.append(title_para)
        elements.append(Spacer(1, 0.2*inch))
        
        # Авторы
        if self.report_data.authors:
            authors_text = "<b>Авторы:</b> " + ", ".join(self.report_data.authors)
            elements.append(Paragraph(authors_text, styles['Normal']))
            elements.append(Spacer(1, 0.1*inch))
        
        # Журнал и дата
        journal_text = f"<b>Журнал:</b> {self.report_data.journal}"
        elements.append(Paragraph(journal_text, styles['Normal']))
        
        if self.report_data.publication_date and self.report_data.publication_date != "N/A":
            pub_date = self.report_data.publication_date[:10] if len(str(self.report_data.publication_date)) > 10 else self.report_data.publication_date
            elements.append(Paragraph(f"<b>Дата публикации:</b> {pub_date}", styles['Normal']))
        elements.append(Spacer(1, 0.1*inch))
        
        # DOI
        if self.report_data.doi and self.report_data.doi != "N/A":
            elements.append(Paragraph(f"<b>DOI:</b> {self.report_data.doi}", styles['Normal']))
            elements.append(Spacer(1, 0.2*inch))
        
        # Аннотация
        elements.append(Paragraph("<b>Аннотация</b>", styles['Heading3']))
        if self.report_data.abstract:
            abstract_text = self.report_data.abstract[:2000] + "..." if len(self.report_data.abstract) > 2000 else self.report_data.abstract
            elements.append(Paragraph(abstract_text, styles['Normal']))
        else:
            elements.append(Paragraph("<i>Аннотация отсутствует</i>", styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        
        # Ключевые слова
        elements.append(Paragraph("<b>Ключевые слова</b>", styles['Heading3']))
        if self.report_data.keywords:
            keywords_text = ", ".join(self.report_data.keywords)
            elements.append(Paragraph(keywords_text, styles['Normal']))
        else:
            elements.append(Paragraph("<i>Ключевые слова не указаны</i>", styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        
        # Метрики
        elements.append(Paragraph("<b>Метрики</b>", styles['Heading3']))
        
        metrics_data = [
            ["Параметр", "Значение"],
            ["Длина аннотации", f"{self.report_data.abstract_length} символов"],
            ["Длина полного текста", f"{self.report_data.full_text_length} символов"],
            ["Количество ключевых слов", str(self.report_data.keywords_count)],
            ["Оценка качества", f"{self.report_data.quality_score}/100"],
        ]
        
        metrics_table = Table(metrics_data, colWidths=[3*inch, 2*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a6cf7')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f1f5f9')]),
        ]))
        elements.append(metrics_table)
        elements.append(Spacer(1, 0.3*inch))
        
        # Рекомендации
        recommendations = self.report_data.get_recommendations()
        elements.append(Paragraph("<b>Рекомендации</b>", styles['Heading3']))
        
        if recommendations:
            for rec in recommendations:
                elements.append(Paragraph(f"• {rec}", styles['Normal']))
        else:
            elements.append(Paragraph("<i>Нет рекомендаций. Статья соответствует критериям качества.</i>", styles['Normal']))
        
        # Построение PDF
        doc.build(elements)
        
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
    
    def export_docx(self) -> bytes:
        """
        Экспортировать отчёт в DOCX.
        
        Returns:
            Bytes DOCX файла
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx не установлен. Установите: pip install python-docx")
        
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('ОТЧЁТ ПО СТАТЬЕ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Основная информация
        doc.add_paragraph(f"ID: {self.report_data.id}", style='Intense Quote')
        doc.add_paragraph(f"Источник: {self.report_data.source}")
        doc.add_paragraph(f"Дата генерации: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        # Название
        doc.add_heading(self.report_data.title, level=1)
        
        # Авторы
        if self.report_data.authors:
            doc.add_paragraph(f"Авторы: {', '.join(self.report_data.authors)}")
        
        # Журнал и дата
        doc.add_paragraph(f"Журнал: {self.report_data.journal}")
        
        if self.report_data.publication_date and self.report_data.publication_date != "N/A":
            pub_date = self.report_data.publication_date[:10] if len(str(self.report_data.publication_date)) > 10 else self.report_data.publication_date
            doc.add_paragraph(f"Дата публикации: {pub_date}")
        
        # DOI
        if self.report_data.doi and self.report_data.doi != "N/A":
            doc.add_paragraph(f"DOI: {self.report_data.doi}")
        
        # Аннотация
        doc.add_heading('Аннотация', level=2)
        if self.report_data.abstract:
            doc.add_paragraph(self.report_data.abstract)
        else:
            doc.add_paragraph('Аннотация отсутствует', style='Intense Quote')
        
        # Ключевые слова
        doc.add_heading('Ключевые слова', level=2)
        if self.report_data.keywords:
            doc.add_paragraph(', '.join(self.report_data.keywords))
        else:
            doc.add_paragraph('Ключевые слова не указаны', style='Intense Quote')
        
        # Метрики
        doc.add_heading('Метрики', level=2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Заголовок таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'
        
        # Жирный шрифт для заголовка
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Данные
        metrics = [
            ("Длина аннотации", f"{self.report_data.abstract_length} символов"),
            ("Длина полного текста", f"{self.report_data.full_text_length} символов"),
            ("Количество ключевых слов", str(self.report_data.keywords_count)),
            ("Оценка качества", f"{self.report_data.quality_score}/100"),
        ]
        
        for param, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = value
        
        # Рекомендации
        doc.add_heading('Рекомендации', level=2)
        
        recommendations = self.report_data.get_recommendations()
        if recommendations:
            for rec in recommendations:
                doc.add_paragraph(rec, style='List Bullet')
        else:
            doc.add_paragraph('Нет рекомендаций. Статья соответствует критериям качества.', style='Intense Quote')
        
        # Сохранение в bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        return docx_bytes


def generate_paper_pdf(paper: dict) -> bytes:
    """
    Сгенерировать PDF отчёт по статье.
    
    Args:
        paper: Данные статьи
        
    Returns:
        Bytes PDF файла
    """
    report_data = PaperReportData(paper)
    exporter = ReportExporter(report_data)
    return exporter.export_pdf()


def generate_paper_docx(paper: dict) -> bytes:
    """
    Сгенерировать DOCX отчёт по статье.
    
    Args:
        paper: Данные статьи
        
    Returns:
        Bytes DOCX файла
    """
    report_data = PaperReportData(paper)
    exporter = ReportExporter(report_data)
    return exporter.export_docx()


if __name__ == "__main__":
    # Пример использования
    sample_paper = {
        "id": 1,
        "title": "Nickel-based superalloys for high-temperature applications",
        "authors": ["John Smith", "Jane Doe"],
        "publication_date": "2024-01-15",
        "journal": "Materials Science and Engineering",
        "doi": "10.1234/test.2024.001",
        "source": "CORE",
        "abstract": "This paper presents a comprehensive review of nickel-based superalloys...",
        "full_text": "Full text content here...",
        "keywords": ["nickel", "superalloys", "high-temperature", "materials"],
    }
    
    # Генерация PDF
    pdf_bytes = generate_paper_pdf(sample_paper)
    print(f"PDF generated: {len(pdf_bytes)} bytes")
    
    # Генерация DOCX
    docx_bytes = generate_paper_docx(sample_paper)
    print(f"DOCX generated: {len(docx_bytes)} bytes")
